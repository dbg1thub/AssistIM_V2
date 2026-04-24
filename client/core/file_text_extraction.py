"""Helpers for local text extraction and summary metadata on file messages."""

from __future__ import annotations

import asyncio
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


FILE_TEXT_EXTRACT_EXTRA_KEY = "file_text_extract"
FILE_SUMMARY_EXTRA_KEY = "file_summary"
FILE_TEXT_EXTRACT_MAX_BYTES = 10 * 1024 * 1024
FILE_TEXT_EXTRACT_MAX_CHARS = 8000
FILE_TEXT_EXTRACT_MAX_PDF_PAGES = 20
FILE_SUMMARY_MAX_CHARS = 720

SUPPORTED_TEXT_FILE_EXTENSIONS = {".txt", ".md", ".markdown"}
SUPPORTED_FILE_TEXT_EXTENSIONS = SUPPORTED_TEXT_FILE_EXTENSIONS | {".pdf", ".docx"}


@dataclass(slots=True)
class FileTextExtractionConfig:
    """Limits for one local file text extraction pass."""

    max_file_bytes: int = FILE_TEXT_EXTRACT_MAX_BYTES
    max_text_chars: int = FILE_TEXT_EXTRACT_MAX_CHARS
    max_pdf_pages: int = FILE_TEXT_EXTRACT_MAX_PDF_PAGES


@dataclass(slots=True)
class FileTextExtractionResult:
    """One local file text extraction result."""

    text: str
    file_name: str = ""
    file_ext: str = ""
    size_bytes: int = 0
    truncated: bool = False
    page_count: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


class FileTextExtractionError(RuntimeError):
    """Stable local file text extraction error."""

    def __init__(self, code: str, message: str = "") -> None:
        self.code = code
        super().__init__(message or code)


class LocalFileTextExtractor:
    """Extract bounded plain text from supported local file attachments."""

    def __init__(self, config: FileTextExtractionConfig | None = None) -> None:
        self._config = config or FileTextExtractionConfig()

    async def extract(self, file_path: str, *, display_name: str = "", mime_type: str = "") -> FileTextExtractionResult:
        del mime_type
        return await asyncio.to_thread(self.extract_sync, file_path, display_name=display_name)

    def extract_sync(self, file_path: str, *, display_name: str = "", mime_type: str = "") -> FileTextExtractionResult:
        del mime_type
        path = self._assert_file_available(file_path)
        file_name = str(display_name or path.name or "").strip() or path.name
        file_ext = Path(file_name).suffix.lower() or path.suffix.lower()
        if file_ext not in SUPPORTED_FILE_TEXT_EXTENSIONS:
            raise FileTextExtractionError("FILE_TEXT_UNSUPPORTED_TYPE", f"Unsupported file type: {file_ext}")

        size_bytes = path.stat().st_size
        if size_bytes > max(1, int(self._config.max_file_bytes or FILE_TEXT_EXTRACT_MAX_BYTES)):
            raise FileTextExtractionError("FILE_TEXT_FILE_TOO_LARGE", "File is too large for local text extraction")

        if file_ext in SUPPORTED_TEXT_FILE_EXTENSIONS:
            text = self._read_text_file(path)
            page_count = 0
        elif file_ext == ".pdf":
            text, page_count = self._read_pdf_file(path)
        elif file_ext == ".docx":
            text, page_count = self._read_docx_file(path)
        else:
            raise FileTextExtractionError("FILE_TEXT_UNSUPPORTED_TYPE", f"Unsupported file type: {file_ext}")

        normalized_text = _normalize_extracted_text(text)
        clipped_text, truncated = _clip_text(normalized_text, self._config.max_text_chars)
        if not clipped_text:
            raise FileTextExtractionError("FILE_TEXT_EMPTY", "No readable text found in file")

        return FileTextExtractionResult(
            text=clipped_text,
            file_name=file_name,
            file_ext=file_ext,
            size_bytes=size_bytes,
            truncated=truncated,
            page_count=page_count,
            metadata={
                "engine": "local_file_text",
                "max_text_chars": max(1, int(self._config.max_text_chars or FILE_TEXT_EXTRACT_MAX_CHARS)),
            },
        )

    @staticmethod
    def _assert_file_available(file_path: str) -> Path:
        path = Path(str(file_path or "")).expanduser()
        if not str(path):
            raise FileTextExtractionError("FILE_TEXT_NOT_FOUND", "File path is required")
        try:
            resolved = path.resolve()
        except OSError as exc:
            raise FileTextExtractionError("FILE_TEXT_NOT_FOUND", str(exc)) from exc
        if not resolved.is_file():
            raise FileTextExtractionError("FILE_TEXT_NOT_FOUND", f"File not found: {resolved}")
        return resolved

    @staticmethod
    def _read_text_file(path: Path) -> str:
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")

    def _read_pdf_file(self, path: Path) -> tuple[str, int]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise FileTextExtractionError("FILE_TEXT_DEPENDENCY_MISSING", "pypdf is not installed") from exc

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        if page_count > max(1, int(self._config.max_pdf_pages or FILE_TEXT_EXTRACT_MAX_PDF_PAGES)):
            raise FileTextExtractionError("FILE_TEXT_TOO_MANY_PAGES", "PDF has too many pages")
        parts = [str(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(part for part in parts if part), page_count

    @staticmethod
    def _read_docx_file(path: Path) -> tuple[str, int]:
        try:
            from docx import Document
        except ImportError as exc:
            raise FileTextExtractionError("FILE_TEXT_DEPENDENCY_MISSING", "python-docx is not installed") from exc

        document = Document(str(path))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts), 0


_file_text_extractor: LocalFileTextExtractor | None = None


def get_local_file_text_extractor() -> LocalFileTextExtractor:
    """Return the process-wide local file text extractor."""
    global _file_text_extractor
    if _file_text_extractor is None:
        _file_text_extractor = LocalFileTextExtractor()
    return _file_text_extractor


def file_summary_display_text(extra: dict[str, Any] | None) -> str:
    """Return the default display text for local file summary/extraction metadata."""
    data = dict(extra or {})
    summary = dict(data.get(FILE_SUMMARY_EXTRA_KEY) or {})
    summary_status = str(summary.get("status") or "").strip()
    if summary_status == "ready":
        return str(summary.get("text") or "").strip()
    if summary_status == "pending":
        return "正在总结文件内容..."
    if summary_status == "failed":
        return "文件总结失败"

    extraction = dict(data.get(FILE_TEXT_EXTRACT_EXTRA_KEY) or {})
    extract_status = str(extraction.get("status") or "").strip()
    reason = str(extraction.get("reason") or "").strip()
    if extract_status == "pending":
        return "正在读取文件内容..."
    if extract_status == "skipped" and reason == "unsupported_type":
        return "暂不支持总结该文件类型"
    if extract_status == "skipped" and reason == "file_too_large":
        return "文件过大，暂不支持总结"
    if extract_status == "skipped" and reason == "too_many_pages":
        return "PDF 页数过多，暂不支持总结"
    if extract_status == "failed" and reason == "dependency_missing":
        return "缺少文件解析依赖"
    if extract_status == "failed":
        return "文件内容读取失败"
    return ""


def extracted_file_context_text(extra: dict[str, Any] | None, *, max_chars: int) -> str:
    """Return ready extracted file text for AI/summary context."""
    extraction = dict((extra or {}).get(FILE_TEXT_EXTRACT_EXTRA_KEY) or {})
    if str(extraction.get("status") or "").strip() != "ready":
        return ""
    text = _normalize_extracted_text(extraction.get("text"))
    clipped, _truncated = _clip_text(text, max_chars)
    return clipped


def _normalize_extracted_text(value: Any) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _clip_text(text: str, max_chars: int) -> tuple[str, bool]:
    normalized_max = max(1, int(max_chars or FILE_TEXT_EXTRACT_MAX_CHARS))
    if len(text) <= normalized_max:
        return text, False
    return text[:normalized_max].rstrip(), True
